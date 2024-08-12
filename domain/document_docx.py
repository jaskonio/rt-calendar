import re
from docx import Document
from docx.document import Document as DocumentObj


class DocumentDocx():
    document:DocumentObj

    def __init__(self, document_file_name:str) -> None:
        self.document = Document(document_file_name)

    def get_name_athlete(self):
        try:
            first_paragraphs =  self.document.paragraphs[1].text
            regex = r'ATLETA:\s*([\wáéíóúüñÁÉÍÓÚÜÑ\s]+)\s+SEMANA:'
            name_athlete = re.findall(regex, first_paragraphs, re.IGNORECASE)[0].strip()

            return name_athlete
        except Exception as error:
            print(f"ERROR getting name")
            print(error)
            return None

    def get_start_end_date(self):
        """_summary_

        Returns:
            list[str]: ['29 de Julio', '11 de Agosto', '2024']
        """
        try:
            ## Get Start/End Date
            date_paragraph = ''

            for content in self.document.paragraphs[1].runs:
                if 'del' in content.text.strip():
                    date_paragraph = content.text.strip()

            if date_paragraph == '':
                return 'Data not found in document'

            regex = r'\(del\s(\d{1,2}\sde\s\w+\s)al(\s\d{1,2}\sde\s\w+)(\s\d{4})\)'
            string_dates = re.findall(regex, date_paragraph, re.IGNORECASE)
            start_date = ''
            end_date = ''
            year = ''

            if len(string_dates) != 0:
                string_dates = string_dates[0]
                start_date = string_dates[0].strip()
                end_date = string_dates[1].strip()
                year = string_dates[2].strip()
            else:
                # Only month
                regex = r'\(del\s(\d{1,2})\sal\s(\d{1,2})\sde\s(\w+)\s(\d{4})\)'
                string_dates = re.findall(regex, date_paragraph, re.IGNORECASE)
                string_dates = string_dates[0]

                start_date = f'{string_dates[0]} de {string_dates[2]}'
                end_date = f'{string_dates[1]} de {string_dates[2]}'
                year = str(string_dates[3])

            return [start_date, end_date, year]
        except Exception as error:
            print(f"ERROR get_start_end_date")
            print(error)
            return None

    def get_training(self):
        try:
            n_days = range(0,7)
            index_first_week = 1
            index_second_week = 2

            content_first_week = [self.document.tables[0].cell(index_first_week,day).text for day in n_days]
            content_second_week = [self.document.tables[0].cell(index_second_week, day).text for day in n_days]
            content_trainer = content_first_week + content_second_week

            return content_trainer
        except Exception as error:
            print(f"ERROR get_training")
            print(error)
            return None
